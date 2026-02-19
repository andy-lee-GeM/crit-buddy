"""
Word document generator for calculation reports.

Creates properly formatted .docx files following the calculation template structure.
"""

from pathlib import Path
from typing import Dict, Union, Optional
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_calculation_docx(
    markdown_path: Union[str, Path],
    output_path: Union[str, Path],
    template_path: Optional[Union[str, Path]] = None,
) -> Path:
    """
    Generate a formatted Word document from a markdown calculation report.

    Args:
        markdown_path: Path to the markdown report
        output_path: Path for output docx
        template_path: Optional path to Word template for style reference

    Returns:
        Path to generated docx
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx required. Install with: pip install python-docx")

    markdown_path = Path(markdown_path)
    output_path = Path(output_path)

    # Read markdown content
    with open(markdown_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    _setup_styles(doc)

    # Parse and add content
    _parse_markdown_to_docx(doc, md_content, markdown_path.parent)

    # Save
    doc.save(output_path)

    return output_path


def _setup_styles(doc: Document):
    """Configure document styles to match calculation template."""
    styles = doc.styles

    # Title style
    if 'Title' in [s.name for s in styles]:
        title_style = styles['Title']
        title_style.font.size = Pt(26)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 1 - Main sections (1. References, 2. Purpose, etc.)
    h1 = styles['Heading 1']
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2 - Subsections (3.1, 3.2, etc.)
    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 51, 102)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3 - Sub-subsections
    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 51, 102)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # Normal text
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'
    normal.paragraph_format.space_after = Pt(8)


def _parse_markdown_to_docx(doc: Document, content: str, base_path: Path):
    """Parse markdown content and add to document."""
    lines = content.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            # Add a subtle separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            # Title (H1 in markdown = document title)
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith('## '):
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            text = line[3:].strip()
            doc.add_paragraph(text, style='Heading 1')
            i += 1
            continue

        if line.startswith('### '):
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            text = line[4:].strip()
            doc.add_paragraph(text, style='Heading 2')
            i += 1
            continue

        if line.startswith('#### '):
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            text = line[5:].strip()
            doc.add_paragraph(text, style='Heading 3')
            i += 1
            continue

        # Table row
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Check if it's a separator row
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            # Parse table row
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue

        # Image
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            _add_image(doc, base_path / img_path, alt_text)
            i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            text = line.strip()[2:]
            text = _process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if num_match:
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            text = num_match.group(2)
            text = _process_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
            i += 1
            continue

        # Block quote
        if line.strip().startswith('>'):
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            text = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        # Italic caption (starts with *)
        if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            text = line.strip()[1:-1]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph
        if in_table and table_rows:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        text = _process_inline_formatting(line.strip())
        p = doc.add_paragraph()
        _add_formatted_text(p, text)
        i += 1

    # Handle any remaining table
    if table_rows:
        _add_table(doc, table_rows)


def _process_inline_formatting(text: str) -> str:
    """Process inline markdown formatting markers."""
    return text


def _add_formatted_text(paragraph, text: str):
    """Add text with inline formatting to a paragraph."""
    # Simple pattern matching for bold and inline code
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def _add_table(doc: Document, rows: list):
    """Add a formatted table to the document."""
    if not rows:
        return

    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text

            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # Add space after table
    doc.add_paragraph()


def _add_image(doc: Document, img_path: Path, alt_text: str):
    """Add an image to the document."""
    if img_path.exists():
        try:
            doc.add_picture(str(img_path), width=Inches(6))
            # Add caption
            p = doc.add_paragraph()
            run = p.add_run(f"Figure: {alt_text}")
            run.italic = True
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            # If image fails, add placeholder text
            p = doc.add_paragraph(f"[Image: {alt_text}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Image not found: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
